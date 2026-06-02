#!/usr/bin/env python3
"""Generate a sample with REAL tracked changes for the Revision Processor.

Run:  python make_example.py
Produces:  sample_review.docx  (next to this script)

WHY THIS IS NON-TRIVIAL (the python-docx revision-authoring gap)
----------------------------------------------------------------
python-docx can write paragraphs and runs, but it has NO API for authoring
tracked changes: the OOXML revision elements ``<w:ins>`` (a tracked
INSERTION) and ``<w:del>`` (a tracked DELETION, whose text lives in
``<w:delText>`` rather than ``<w:t>``) cannot be produced through its
high-level model. This is the same class of gap as python-pptx's inability
to author review comments (see ``rags/lessons/pptx_fixture_authoring.md``).

So this generator uses APPROACH (a): build the base document with python-docx,
then re-open the .docx zip and INJECT genuine ``<w:ins>`` / ``<w:del>``
revision markup directly into ``word/document.xml`` by wrapping runs, then
re-zip. This mirrors the OOXML zip-rewrite technique used by the Link
Auditor's ``examples/make_example.py`` (which injects external-link parts).

The result is a genuinely revision-bearing document: opening it in Word shows
tracked changes in the markup, and ``wdDoc.Revisions.Count`` returns > 0 — so
the tool has real revisions to accept or reject.

VERIFICATION: after writing, this script re-opens the .docx zip and asserts
that ``word/document.xml`` contains both ``<w:ins`` and ``<w:del``.

ASCII-only print() (the Bash tool's console is cp1252 and chokes on non-ASCII).
"""
import re
import shutil
import zipfile
from pathlib import Path

import docx

HERE = Path(__file__).resolve().parent
OUT = HERE / "sample_review.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REV_AUTHOR = "Reviewer"
REV_DATE = "2026-01-01T00:00:00Z"


def build_base() -> None:
    """Write the base document. Each paragraph carries a unique marker run so
    we can find it in document.xml and wrap it in revision markup."""
    doc = docx.Document()
    doc.add_heading("Project Proposal", level=1)

    # Paragraph 1: a run we will mark as a tracked INSERTION.
    doc.add_paragraph(
        "The team will deliver the prototype by the end of the quarter. "
        "INS_MARKER_This sentence was added by the reviewer._END "
        "We remain on schedule.")

    # Paragraph 2: a run we will mark as a tracked DELETION.
    doc.add_paragraph(
        "Budget remains within the approved envelope. "
        "DEL_MARKER_This obsolete clause should be removed._END "
        "No further funding is required.")

    # Paragraph 3: plain text, no revisions (proves mixed content survives).
    doc.add_paragraph(
        "Please direct questions to the project office.")

    doc.save(str(OUT))


def inject_revisions() -> None:
    """Re-open the .docx zip and wrap the marker runs in w:ins / w:del.

    python-docx emits each paragraph's text as one or more
    ``<w:r><w:t>...</w:t></w:r>`` runs. We locate the run that contains our
    marker text, split it into before/marked/after, and wrap the marked run
    in the appropriate revision element. For a deletion the inner ``<w:t>``
    becomes ``<w:delText>`` (Word stores deleted text there)."""
    tmp = OUT.with_suffix(".tmp.docx")
    with zipfile.ZipFile(OUT, "r") as zin:
        names = zin.namelist()
        data = {n: zin.read(n) for n in names}

    xml = data["word/document.xml"].decode("utf-8")

    # --- tracked INSERTION -------------------------------------------------
    ins_text = "This sentence was added by the reviewer."
    ins_run = (
        f'<w:ins w:id="901" w:author="{REV_AUTHOR}" w:date="{REV_DATE}">'
        f'<w:r><w:t xml:space="preserve">{ins_text}</w:t></w:r>'
        f'</w:ins>'
    )
    # Our base text put "INS_MARKER_<text>_END" inside one <w:t>. Replace the
    # whole "<w:r>...INS_MARKER_..._END...</w:r>" surrounding run with three
    # runs: plain-before, the w:ins, plain-after. Simplest robust approach:
    # the marker text is unique, so swap the literal marker substring for a
    # run-break + w:ins + run-break that re-opens a plain run.
    xml = xml.replace(
        f"INS_MARKER_{ins_text}_END",
        f'</w:t></w:r>{ins_run}<w:r><w:t xml:space="preserve">',
    )

    # --- tracked DELETION --------------------------------------------------
    del_text = "This obsolete clause should be removed."
    del_run = (
        f'<w:del w:id="902" w:author="{REV_AUTHOR}" w:date="{REV_DATE}">'
        f'<w:r><w:delText xml:space="preserve">{del_text}</w:delText></w:r>'
        f'</w:del>'
    )
    xml = xml.replace(
        f"DEL_MARKER_{del_text}_END",
        f'</w:t></w:r>{del_run}<w:r><w:t xml:space="preserve">',
    )

    data["word/document.xml"] = xml.encode("utf-8")

    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, blob in data.items():
            zout.writestr(name, blob)
    shutil.move(str(tmp), str(OUT))


def verify() -> int:
    """Re-open and confirm both revision elements are present. Returns the
    number of revision elements found (ins + del)."""
    with zipfile.ZipFile(OUT, "r") as z:
        xml = z.read("word/document.xml").decode("utf-8")
    ins = len(re.findall(r"<w:ins[ >]", xml))
    dele = len(re.findall(r"<w:del[ >]", xml))
    print(f"verify: found {ins} <w:ins> and {dele} <w:del> element(s)")
    assert ins >= 1, "no <w:ins> insertion was injected"
    assert dele >= 1, "no <w:del> deletion was injected"
    # also confirm python-docx can still open it (well-formed)
    docx.Document(str(OUT))
    print("verify: python-docx re-opened the document OK (well-formed)")
    return ins + dele


def main() -> None:
    build_base()
    inject_revisions()
    n = verify()
    print(f"wrote {OUT}  ({n} tracked change(s) injected)")


if __name__ == "__main__":
    main()
