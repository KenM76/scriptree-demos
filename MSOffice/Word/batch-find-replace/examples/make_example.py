#!/usr/bin/env python3
"""Generate the sample memo for the Batch Find & Replace tool.

Run:  python make_example.py
Produces:  sample_memo.docx  (next to this script)

WHAT THE SAMPLE CONTAINS (and why)
----------------------------------
The doc is seeded so you can see every matching option do something
different against the SAME find term, ``Foobar``:

* **Five literal "Foobar" occurrences** across a heading and body
  paragraphs → a clean replacement count to verify.
* **One lower-case "foobar"** → flips in/out of the count depending on
  **Match case**.
* **One "Foobars" (plural)** and one **"Foobaristas"** → these contain
  "Foobar" as a substring, so they DROP OUT of the count when **Whole
  words only** is ticked. A great way to see the whole-word guard work.

Counting the bold "Foobar" in the heading, the sample has exactly **5**
standalone "Foobar", **1** "foobar", **1** "Foobars", **1**
"Foobaristas".
"""
from pathlib import Path

from docx import Document


def main() -> None:
    doc = Document()

    h = doc.add_heading("Foobar Project — Status Memo", level=1)
    # add_heading already styled it; the word "Foobar" in the heading counts.

    doc.add_paragraph(
        "This memo summarises the Foobar rollout. The Foobar team met on "
        "Monday and confirmed that Foobar is on schedule."
    )
    doc.add_paragraph(
        "Note: an early draft spelled it foobar in lower case; that "
        "variant is here on purpose to test the Match case option."
    )
    doc.add_paragraph(
        "Several Foobars were shipped last week, and the Foobaristas in "
        "support handled the questions. (These two words merely CONTAIN "
        "'Foobar' and should be excluded when Whole words only is on.)"
    )
    doc.add_paragraph(
        "In closing, Foobar remains our priority for the quarter."
    )

    out = Path(__file__).resolve().parent / "sample_memo.docx"
    doc.save(str(out))
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
