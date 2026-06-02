#!/usr/bin/env python3
"""Generate the sample manuscript for the Heading-Based Document Splitter.

Run:  python make_example.py
Produces:  sample_manuscript.docx  (next to this script)

WHAT THE SAMPLE CONTAINS (and why each piece is here)
-----------------------------------------------------
The splitter cuts at paragraphs whose OUTLINE LEVEL equals the chosen
heading level (it keys off the built-in "Heading N" styles, which set
that outline level). So the sample is built to exercise every branch of
the tool:

* **Front matter** — a title line + two abstract paragraphs that carry
  NO heading style. They sit *before* the first Heading 1, so with
  "Save content before first heading" on they become ``00_Front-Matter``.
* **Four Heading 1 sections** — Introduction, Methods, Results,
  Conclusion. Splitting at level 1 yields one file per section
  (``01_Introduction`` … ``04_Conclusion``).
* **Two Heading 2 sub-sections inside Methods** — "Data Collection" and
  "Analysis". Splitting at level 2 instead yields a file per *sub*-section
  (and the level-1 headings then fall inside whichever level-2 block
  precedes the next level-2 heading), which is exactly how you show the
  finer-grained split.
* **Body paragraphs** under each heading so the resulting files are not
  empty and the FormattedText copy has real content to carry.

The file is intentionally tiny and text-only; the point is to demonstrate
the split boundaries, not to be a realistic 200-page book.
"""
from pathlib import Path

from docx import Document


def main() -> None:
    doc = Document()

    # --- front matter (no heading styles → becomes 00_Front-Matter) -------
    title = doc.add_paragraph("A Short Study of Widget Fatigue")
    title.runs[0].bold = True
    doc.add_paragraph(
        "Abstract. This document is a sample manuscript used to demonstrate "
        "the Heading-Based Document Splitter. The text here precedes the "
        "first heading, so it is the document's front matter."
    )
    doc.add_paragraph(
        "Prepared as example input. Splitting at Heading 1 produces one "
        "file per top-level section; splitting at Heading 2 produces one "
        "file per sub-section within Methods."
    )

    # --- Section 1: Introduction (Heading 1) ------------------------------
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Widgets are everywhere, and yet little is known about how they "
        "tire. This study sets out to change that."
    )

    # --- Section 2: Methods (Heading 1) with two Heading 2 sub-sections ----
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("We combined field observation with bench testing.")

    doc.add_heading("Data Collection", level=2)
    doc.add_paragraph(
        "Widgets were sampled weekly over a three-month period and logged "
        "by serial number."
    )

    doc.add_heading("Analysis", level=2)
    doc.add_paragraph(
        "Fatigue was scored on a five-point scale and aggregated per cohort."
    )

    # --- Section 3: Results (Heading 1) -----------------------------------
    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "Fatigue rose sharply after the eighth week of continuous use."
    )

    # --- Section 4: Conclusion (Heading 1) --------------------------------
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "Widget fatigue is real, measurable, and worth designing around."
    )

    out = Path(__file__).resolve().parent / "sample_manuscript.docx"
    doc.save(str(out))
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
