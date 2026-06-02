#!/usr/bin/env python3
"""Generate the deliberately-messy sample for the Corporate Style Sanitizer.

Run:  python make_example.py
Produces:  sample_messy.docx  (next to this script)

WHAT THE SAMPLE CONTAINS (and why)
----------------------------------
The sanitizer strips manual formatting, tidies whitespace, normalises fonts,
and (optionally) curls quotes onto a sibling ``_Sanitized`` copy. The sample
is seeded so every pass has something to do AND so you can confirm the things
that must be PRESERVED survive:

* **Rogue direct formatting** — headings and body text carry manual font/size/
  colour/bold overrides on top of their styles (Comic Sans, red, 20pt, etc.).
  "Strip manual formatting" should revert these to the style.
* **Messy whitespace** — double (and triple) spaces inside sentences and
  trailing spaces before paragraph marks. "Collapse repeated & trailing
  spaces" should fix them.
* **Blank spacer paragraphs** — several empty paragraphs hammered in for
  vertical spacing. "Remove blank paragraphs" should collapse them.
* **Straight quotes** — 'single' and "double" straight quotes. "Convert
  straight quotes to curly" (off by default) should curl them.
* **Real headings via Heading 1/2 styles** — so "Normalise fonts" has heading
  styles to rewrite.

PRESERVED (must survive every pass):
* **A real hyperlink** — the field/URL must remain clickable; only its manual
  underline/colour reverts when formatting is stripped.
* **A shaded table** — the cell background fill must remain (paragraph-format
  reset does not touch cell shading).
"""
from pathlib import Path

import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def add_hyperlink(paragraph, url, text):
    """Insert a real hyperlink (relationship + w:hyperlink) into a paragraph.

    python-docx has no high-level hyperlink API, so we add the relationship on
    the document part and build the <w:hyperlink><w:r>...</w:r></w:hyperlink>
    by hand. This is the standard recipe; the link is a genuine field so the
    sanitizer can be shown to preserve it.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = docx.oxml.OxmlElement("w:r")
    rpr = docx.oxml.OxmlElement("w:rPr")
    # Manual link styling (blue + underline) so we can see it revert on strip.
    color = docx.oxml.OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = docx.oxml.OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(u)
    run.append(rpr)
    t = docx.oxml.OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def shade_cell(cell, fill_hex):
    """Set a table cell's background fill (w:shd) — should survive sanitising."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = docx.oxml.OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def rogue_run(paragraph, text, *, font=None, size=None, color=None,
              bold=False, highlight=False):
    """Add a run with deliberate manual (direct) formatting overrides."""
    r = paragraph.add_run(text)
    if font:
        r.font.name = font
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    r.font.bold = bold
    if highlight:
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return r


def main() -> None:
    doc = docx.Document()

    # --- Heading 1 with a rogue manual override -------------------------
    h1 = doc.add_heading("", level=1)
    rogue_run(h1, "Quarterly  Business  Review", font="Comic Sans MS",
              color=(0xC0, 0x00, 0x00), size=22)

    # body paragraph: double spaces + trailing spaces + straight quotes +
    # manual font/colour
    p = doc.add_paragraph()
    rogue_run(
        p,
        'This  document  was  pasted  together  from  many  sources.   ',
        font="Arial", size=13, color=(0x20, 0x60, 0xA0))
    rogue_run(
        p,
        'The boss said "make it look consistent" and \'clean it up\'.   ',
        font="Georgia", size=13, bold=True)

    # blank spacer paragraphs (hammered Enter)
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # --- Heading 2 ------------------------------------------------------
    h2 = doc.add_heading("", level=2)
    rogue_run(h2, "Key  Links", font="Times New Roman", color=(0x00, 0x80, 0x00))

    # paragraph with a REAL hyperlink (must be preserved)
    p2 = doc.add_paragraph()
    p2.add_run("Reference site: ")
    add_hyperlink(p2, "https://example.com/standards", "Corporate Standards")
    p2.add_run("  (keep this link working).")

    doc.add_paragraph("")
    doc.add_paragraph("")

    # --- Heading 2 + a shaded table (shading must be preserved) ---------
    h2b = doc.add_heading("", level=2)
    rogue_run(h2b, "Summary  Table", font="Comic Sans MS", size=18)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    cells = [("Region", "Revenue"), ("East", "1,200")]
    for r in range(2):
        for c in range(2):
            cell = table.cell(r, c)
            cell.text = cells[r][c]
            # shade the header row blue, body row light grey
            shade_cell(cell, "D9E2F3" if r == 0 else "F2F2F2")

    doc.add_paragraph("")

    # --- a bulleted list (for the ADVANCED 'Normalise bullet glyphs' pass) ---
    # python-docx applies the built-in "List Bullet" style (standard bullets).
    # To see a NON-standard bullet normalised, change one of these to a
    # Wingdings arrow by hand in Word (Home > Bullets > Define New Bullet >
    # Symbol) and re-run with "Normalise bullet glyphs" on.
    h2c = doc.add_heading("", level=2)
    rogue_run(h2c, "Action  Items", font="Comic Sans MS")
    for item in ("Review the draft", "Circulate for sign-off", "Publish"):
        doc.add_paragraph(item, style="List Bullet")

    p3 = doc.add_paragraph()
    rogue_run(
        p3,
        'Footnote:  figures  are  "unaudited"  and  may  change.',
        font="Courier New", size=11, color=(0x80, 0x80, 0x80))

    # --- non-standard page margins (for the ADVANCED 'Enforce margins' pass) -
    # 0.5" all round, so enforcing 1" visibly changes the layout.
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    out = Path(__file__).resolve().parent / "sample_messy.docx"
    doc.save(str(out))
    print(f"wrote {out}")


def make_template() -> None:
    """A clean reference doc with distinct STYLE definitions, for the
    'Style template document' option. Point the sanitizer's style_template at
    this file and sample_messy.docx adopts these styles (Georgia body, navy
    Cambria headings)."""
    tpl = docx.Document()

    normal = tpl.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    for lvl, size in ((1, 18), (2, 14)):
        h = tpl.styles[f"Heading {lvl}"]
        h.font.name = "Cambria"
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)  # navy

    # A little real content so it's an ordinary document, not just a stylesheet.
    tpl.add_heading("Corporate Template", level=1)
    tpl.add_paragraph("Body text uses Georgia 11pt; headings use navy Cambria.")
    tpl.add_heading("Section", level=2)
    tpl.add_paragraph("Point the sanitizer's Style template field at this file.")

    out = Path(__file__).resolve().parent / "sample_template.docx"
    tpl.save(str(out))
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
    make_template()
